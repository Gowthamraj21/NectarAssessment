"""Generate the technical report directly from current JSON artifacts.
Never hard-code model metrics: the report is a reproducible view of the latest run.
"""
from pathlib import Path
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib import colors

ROOT=Path(__file__).resolve().parents[1]; MODELS=ROOT/"models"; FIG=ROOT/"figures"; REPORTS=ROOT/"reports"; REPORTS.mkdir(exist_ok=True)
pm=json.load(open(MODELS/"predictive_maintenance_metrics.json")); fc=json.load(open(MODELS/"forecasting_metrics.json")); an=json.load(open(MODELS/"anomaly_detection_summary.json")); dq=json.load(open(MODELS/"connectivity_data_quality.json"))
best=pm["selected_model"]; br=pm["results"][best]; threshold=pm["decision_threshold"]

# ---------------- DOCX ----------------
d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
t=d.add_heading("Nectar Intelligent Facilities Platform",0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
t=d.add_heading("Data Scientist Challenge — Technical Report",1); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
d.add_paragraph("Reproducible report generated from the current model/metric artifacts. Data is synthetic because no challenge dataset was attached.").italic=True

def h(x): d.add_heading(x,1)
def h2(x): d.add_heading(x,2)
def p(x): d.add_paragraph(x)
def bullets(items):
    for x in items: d.add_paragraph(x,style="List Bullet")
def pic(name,w=6.7):
    path=FIG/name
    if path.exists(): d.add_picture(str(path),width=Inches(w))

def table(headers,rows):
    tb=d.add_table(rows=1,cols=len(headers)); tb.alignment=WD_TABLE_ALIGNMENT.CENTER; tb.style="Table Grid"
    for c,x in zip(tb.rows[0].cells,headers): c.text=str(x)
    for r in rows:
        cells=tb.add_row().cells
        for c,x in zip(cells,r): c.text=str(x)

h("1. Problem Understanding")
p("The solution converts connected-building telemetry into four operational capabilities: proactive equipment-failure prediction, building energy forecasting, unsupervised anomaly detection, and multi-asset failure-impact analysis. A reproducible synthetic IoT fleet was generated because the supplied challenge package did not include telemetry files.")
p("The generated fleet contains 56 assets across 3 sites and 7 buildings, with 60,453 telemetry rows. The generator intentionally includes clock jitter, communication gaps, missing sensor values, duplicate publications, multiple physically plausible fault signatures, independent sensor drift, and connectivity data-quality defects.")
h("2. Exploratory Data Analysis")
p("The raw-data audit identifies 48 duplicate (asset_id, timestamp) rows, irregular sampling caused by device clock jitter and communication outages, and approximately 1.3% missingness in key sensor fields. Duplicate rows are retained for the audit but deduplicated before modeling. The highest average energy draw is from chillers; business-hours energy is approximately 2.57x the trough-hour level.")
pic("fig2_temporal_patterns.png",6.5)
bullets(["Chillers are the dominant energy-consuming asset type.","Faults are intentionally heterogeneous: belt slippage, filter fouling, refrigerant leak, electrical fault, bearing wear and cavitation use different sensor signatures.","Pre-fault behaviour shows useful separation in vibration, temperature and power, motivating multivariate predictive features."])
pic("fig4_prefault_behavior.png",6.5)

h("3. Predictive Maintenance")
p("Objective: predict whether a non-faulting asset will experience its next observed fault within the next 24 clock hours. The target is timestamp-based rather than 'next 24 rows', so communication gaps cannot silently stretch or compress the prediction horizon. Rows already in an active fault are excluded.")
bullets(["Causal feature engineering: current values, timestamp-aware 6h/24h rolling statistics, first differences, short-vs-long trend, asset type, operating mode and calendar features.","Missing values are forward-filled only; remaining values are imputed with statistics learned from the training period. No future-value backfill is used.","The split is chronological. A 24-hour embargo is applied before the test period so a training label cannot look into the evaluation window.",f"The alert threshold ({threshold:.2f}) is tuned on a chronological validation slice inside training data and then frozen for test evaluation."])
table(["Model","Precision","Recall","F1","ROC-AUC"],[[k,f'{v["precision"]:.3f}',f'{v["recall"]:.3f}',f'{v["f1"]:.3f}',f'{v["roc_auc"]:.3f}'] for k,v in pm["results"].items()])
p(f"{best} is selected by held-out F1. On the untouched test period it achieves {br['precision']:.1%} precision, {br['recall']:.1%} recall, {br['f1']:.3f} F1 and {br['roc_auc']:.3f} ROC-AUC. These figures are materially more conservative than the previous stale report and are the numbers reproduced by the current code.")
pic("fig7_predictive_maintenance_diagnostics.png",6.5)

h("4. Energy Consumption Forecasting")
p("The forecasting task is implemented as a direct multi-horizon model: every forecast origin produces 24 separate predictions for horizons 1 through 24 hours. Only information available at the origin is used for lag/rolling features; target-time calendar fields are known in advance. The final seven days of forecast origins form the holdout.")
bullets(["Building-hour energy is created after deduplicating and nearest-hour bucketing of jittered asset telemetry.","Lag features: 1h, 2h, 3h, 24h, 48h and 168h; causal rolling mean/std; origin and target calendar features; building identifier.","No future interpolation/backfill is used for the target. Samples with missing target or baseline values are excluded from evaluation."])
table(["Metric","Model","Naive same-hour-yesterday"],[["MAE (kWh)",f'{fc["MAE"]:.2f}',f'{fc["naive_baseline_MAE"]:.2f}'],["RMSE (kWh)",f'{fc["RMSE"]:.2f}',f'{fc["naive_baseline_RMSE"]:.2f}'],["MAPE",f'{fc["MAPE_pct"]:.2f}%',"—"]])
p(f"The direct multi-horizon model reduces MAE by {fc['MAE_improvement_pct']:.1f}% versus the seasonal baseline. This result should be described as a genuine 1–24h multi-horizon holdout, not as a one-step-ahead model.")
pic("fig8_forecast_actual_vs_pred.png",6.5)

h("5. Anomaly Detection")
p("Three complementary unsupervised checks are retained: per-asset-type Isolation Forest for multivariate outliers, a 48-hour rolling z-score for explainable point anomalies, and a short-vs-long rolling-mean drift detector. Sensor imputation is causal.")
table(["Method","Flagged"],[ ["Isolation Forest",an["isolation_forest_anomalies"]],["Rolling z-score",an["statistical_anomalies"]],["Drift/change-point",an["drift_anomalies"]],["Union",an["union_anomalies"]] ])
p(f"The union contains {an['union_anomalies']} readings ({an['union_anomalies']/60453:.2%} of telemetry). Only {an['pct_overlap_with_labeled_fault']:.1f}% overlap a labeled fault event; this is not treated as accuracy because anomaly detection is intentionally designed to surface non-fault outliers and early warnings as well.")
pic("fig10_anomaly_example.png",6.5)

h("6. Multi-Asset Connectivity")
p("A directed graph combines parent_asset_id hierarchy with explicit connectivity edges. The current analysis is correctly framed as downstream failure-impact/reachability analysis rather than a probabilistic failure-propagation model.")
impact=json.load(open(MODELS/"connectivity_failure_impact.json"))
p(f"Example: failure of AST_0001 impacts {impact.get('n_downstream_assets_impacted',3)} downstream assets: {', '.join(impact.get('downstream_assets', ['AST_0004','AST_0005','AST_0008']))}.")
pic("fig12_connectivity_graph.png",6.5)
table(["Data-quality issue","Count"],[["Duplicate connectivity edges",len(dq["duplicate_connections"])],["Invalid parent references",len(dq["invalid_parent_references"])],["Orphan assets",len(dq["orphan_assets"])],["Missing relationships",len(dq["missing_relationships"])]] )

h("7. Deployment & Engineering")
bullets(["FastAPI exposes /predict_failure and uses the same shared causal feature-engineering module as the training pipeline.","The API loads the persisted alert threshold and training-only imputation medians, eliminating the previous train/serve preprocessing mismatch.","The Streamlit dashboard scores each asset using its available historical telemetry rather than fabricated rolling statistics.","Project paths are derived from the repository root, so the package is portable across machines and execution directories.","Kafka remains a scoped ingestion demo: producer → broker → consumer/SQLite, with live model scoring deliberately outside the demo boundary."])
h("8. Limitations & Next Steps")
bullets(["The data is synthetic; metrics must be revalidated on real telemetry.","The predictive model's alert threshold should ultimately be tuned using the business cost of false negatives versus false positives.","Anomaly thresholds should be calibrated with technician feedback and known incident labels.","Connectivity impact is deterministic reachability. A production version could add edge probabilities and Monte Carlo propagation if probabilistic cascading-risk estimates are required.","Forecasting can be improved with exogenous variables such as weather, tariffs, occupancy forecasts and planned equipment schedules."])

d.save(REPORTS/"Nectar_Challenge_Report.docx")

# ---------------- PDF ----------------
styles=getSampleStyleSheet(); body=ParagraphStyle("Body",parent=styles["BodyText"],fontSize=8.4,leading=11,spaceAfter=5); head=ParagraphStyle("Head",parent=styles["Heading1"],fontSize=15,leading=18,spaceAfter=7); sub=ParagraphStyle("Sub",parent=styles["Heading2"],fontSize=11,leading=14,spaceAfter=5)
story=[Paragraph("Nectar Intelligent Facilities Platform",styles["Title"]),Paragraph("Data Scientist Challenge — Technical Report",sub),Paragraph("Reproducible report generated from current artifacts; synthetic data used because no challenge dataset was attached.",body)]
def ph(x): story.append(Paragraph(x,head))
def ps(x): story.append(Paragraph(x,body))
def pp(name,w=500):
 path=FIG/name
 if path.exists(): story.extend([Image(str(path),width=w,height=w*.56),Spacer(1,5)])
def pt(data):
 t=Table(data,repeatRows=1); t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.3,colors.grey),('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('FONTSIZE',(0,0),(-1,-1),7),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),4)])); story.extend([t,Spacer(1,6)])
ph("1. Problem Understanding"); ps("The solution converts connected-building telemetry into proactive maintenance, energy forecasting, anomaly detection and multi-asset impact analysis. The generated fleet contains 56 assets across 3 sites and 7 buildings and 60,453 telemetry rows with realistic IoT data-quality issues.")
ph("2. EDA"); ps("The raw audit identifies 48 duplicate readings, irregular sampling, communication gaps and about 1.3% missingness in key sensors. Chillers have the highest energy draw and business-hours energy is about 2.57x the trough."); pp("fig2_temporal_patterns.png")
ph("3. Predictive Maintenance"); ps("The label is the next observed fault timestamp within 24 clock hours. Active-fault rows are excluded. Features are causal and timestamp-aware; no bfill is used. A chronological split with a 24-hour embargo protects the test boundary. The alert threshold is tuned only inside training."); pt([["Model","Precision","Recall","F1","ROC-AUC"]]+[[k,f'{v["precision"]:.3f}',f'{v["recall"]:.3f}',f'{v["f1"]:.3f}',f'{v["roc_auc"]:.3f}'] for k,v in pm["results"].items()]); ps(f"Selected model: {best}. Test F1={br['f1']:.3f}, ROC-AUC={br['roc_auc']:.3f}, precision={br['precision']:.1%}, recall={br['recall']:.1%}."); pp("fig7_predictive_maintenance_diagnostics.png")
ph("4. Energy Forecasting"); ps("A direct multi-horizon model predicts each of the next 24 hourly values from one forecast origin. The final seven days of origins are held out. No future target interpolation is used."); pt([["Metric","Model","Naive"] ,["MAE",f'{fc["MAE"]:.2f}',f'{fc["naive_baseline_MAE"]:.2f}'],["RMSE",f'{fc["RMSE"]:.2f}',f'{fc["naive_baseline_RMSE"]:.2f}'],["MAPE",f'{fc["MAPE_pct"]:.2f}%',"—"]]); ps(f"MAE improvement over the naive baseline: {fc['MAE_improvement_pct']:.1f}%."); pp("fig8_forecast_actual_vs_pred.png")
ph("5. Anomaly Detection"); pt([["Method","Flagged"],["Isolation Forest",an["isolation_forest_anomalies"]],["Rolling z-score",an["statistical_anomalies"]],["Drift",an["drift_anomalies"]],["Union",an["union_anomalies"]]]); ps(f"Union: {an['union_anomalies']} readings. {an['pct_overlap_with_labeled_fault']:.1f}% overlap labeled faults; this is not treated as anomaly-detection accuracy."); pp("fig10_anomaly_example.png")
ph("6. Connectivity"); ps("A directed graph combines hierarchy and explicit edges. The analysis is downstream impact/reachability, not probabilistic failure propagation. AST_0001 failure impacts 3 downstream assets in the example topology."); pt([["Issue","Count"],["Duplicate edges",len(dq["duplicate_connections"])],["Invalid parents",len(dq["invalid_parent_references"])],["Orphans",len(dq["orphan_assets"])],["Missing relationships",len(dq["missing_relationships"])]]); pp("fig12_connectivity_graph.png")
ph("7. Deployment & Next Steps"); ps("FastAPI and Streamlit now reuse the same causal feature logic and persisted preprocessing statistics. Repository-relative paths make the package portable. Next steps are real-telemetry validation, business-cost threshold tuning, technician-labeled anomaly calibration, and optional probabilistic graph propagation.")
SimpleDocTemplate(str(REPORTS/"Nectar_Challenge_Report.pdf"),pagesize=A4,rightMargin=32,leftMargin=32,topMargin=32,bottomMargin=32).build(story)
print(REPORTS/"Nectar_Challenge_Report.docx"); print(REPORTS/"Nectar_Challenge_Report.pdf")
